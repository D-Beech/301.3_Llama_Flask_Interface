import docx

# Extract text from docx file

def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def extract_text_from_docx_limited(docx_file, max_words=1000):
    doc = docx.Document(docx_file)
    words = []
    for paragraph in doc.paragraphs:
        p_words = paragraph.text.split()
        if len(words) + len(p_words) > max_words:
            words.extend(p_words[:max_words - len(words)])
            break
        words.extend(p_words)
    return ' '.join(words)